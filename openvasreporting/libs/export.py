# -*- coding: utf-8 -*-
#
"""

This File contains main structure of the OpenVAS Reporting tool.
All exports type are defined here.

"""
#
# Project name: OpenVAS Reporting: A tool to convert OpenVAS XML reports into Excel files.
# Project URL: https://github.com/groupecnpp/OpenvasReporting

from re import sub, search

from collections import Counter
from typing import Callable

from .config import Config
from .parsed_data import ResultTree, Host, Vulnerability


def implemented_exporters() -> dict[str, Callable]:
    """
    Enum-link instance containing references to already implemented exporter function

    > implemented_exporters()[key](param[s])
    
    key is a concatenation of the report-type arg + '-' + format arg

    :return: Pointer to exporter function
    """
    return {
        'vulnerability-xlsx': export_to_excel_by_vuln,
        'vulnerability-docx': export_to_word_by_vuln,
        'vulnerability-csv': export_to_csv_by_vuln,
        'host-xlsx': export_to_excel_by_host,
        'host-csv': export_to_csv_by_host,
        'summary-csv': export_summary_to_csv
    }

def _get_collections(vuln_info:list[Vulnerability]) -> tuple[list[Vulnerability], Counter[str], Counter[str], Counter[str]]:
    """
    Sort vulnerability list info according to CVSS (desc) and Name (asc).
    Provide collections to be used in export.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :return: vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family
    :rtype vuln_info: list(Vulnerability)
    :rtype vuln_levels: Counter
    :rtype vuln_host_by_level: Counter
    :rtype vuln_by_family: Counter
    """
    vuln_info.sort(key=lambda key: key.name)
    vuln_info.sort(key=lambda key: key.cvss, reverse=True)
    vuln_levels:Counter[str] = Counter()
    vuln_host_by_level:Counter[str] = Counter()
    vuln_by_family:Counter[str] = Counter()
    # collect host names
    vuln_hostcount_by_level:list[list[str]] =[[] for _ in range(5)]
    level_choices = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3, 'none': 4}

    for _, vuln in enumerate(vuln_info, 1):
        vuln_levels[vuln.level.lower()] += 1
        # add host names to list so we count unquie hosts per level
        level_index:int = level_choices[vuln.level.lower()]

        for _, (host, _) in enumerate(vuln.hosts, 1):
            if host.ip not in vuln_hostcount_by_level[level_index]:
                vuln_hostcount_by_level[level_index].append(host.ip)

        vuln_by_family[vuln.family] += 1

    # now count hosts per level and return
    for level in Config.levels().values():
        vuln_host_by_level[level] = len([*set(vuln_hostcount_by_level[level_choices[level.lower()]])]) # Put host in a set to avoid doublon

    return vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family


def export_to_excel_by_vuln(vuln_info:list[Vulnerability], threat_type_list:list[str], template=None, output_file:str='openvas_report.xlsx') -> None:
    """
    Export vulnerabilities info in an Excel file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param threat_type_list: Type of Threat present
    :type threat_type_list: list(str)

    :param threat_type_list: Type of Threat present
    :type threat_type_list: list(str)

    :param template: Not supported in xlsx-output
    :type template: NoneType

    :param output_file: Filename of the Excel file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import xlsxwriter

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
            
    if not isinstance(threat_type_list, list):
        raise TypeError("Expected list, got '{}' instead".format(type(threat_type_list)))
    else:
        for x in threat_type_list:
            if not isinstance(x, str):
                raise TypeError("Expected str, got '{}' instead".format(type(x)))
            
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    # if template is not None:
    #     raise NotImplementedError("Use of template is not supported in XSLX-output.")

    vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family = _get_collections(vuln_info)

    # ====================
    # FUNCTIONS
    # ====================
    def __row_height(text:str, width:int) -> int:
        return (max((len(text) // width), text.count('\n')) + 1) * 15

    workbook = xlsxwriter.Workbook(output_file)

    workbook.set_properties({
        'title': output_file,
        'subject': 'OpenVAS report',
        'author': 'CNPP',
        'category': 'report',
        'keywords': 'OpenVAS, report',
        'comments': 'Forked from : TheGroundZero (https://github.com/TheGroundZero)'})

    # ====================
    # FORMATTING
    # ====================
    workbook.formats[0].set_font_name('Tahoma')

    format_sheet_title_content = workbook.add_format({'font_name': 'Tahoma', 'font_size': 12,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'center', 'valign': 'vcenter', 'border': 1})
    format_table_titles = workbook.add_format({'font_name': 'Tahoma', 'font_size': 11,
                                               'font_color': 'white', 'bold': True,
                                               'align': 'center', 'valign': 'vcenter',
                                               'border': 1,
                                               'bg_color': Config.colors()['blue']})
    format_table_cells = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                              'align': 'left', 'valign': 'top',
                                              'border': 1, 'text_wrap': 1})
    format_align_center = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top'})
    format_align_border = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_toc = {
        'critical': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                         'align': 'center', 'valign': 'top',
                                         'border': 1,
                                         'bg_color': Config.colors()['critical']}),
        'high': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['high']}),
        'medium': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                       'align': 'center', 'valign': 'top',
                                       'border': 1, 'bg_color': Config.colors()['medium']}),
        'low': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                    'align': 'center', 'valign': 'top',
                                    'border': 1, 'bg_color': Config.colors()['low']}),
        'none': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['none']})
    }

    # ====================
    # SUMMARY SHEET
    # ====================
    sheet_name = "Summary"
    ws_sum = workbook.add_worksheet(sheet_name)
    ws_sum.set_tab_color(Config.colors()['blue'])

    ws_sum.set_column("A:A", 7, format_align_center)
    ws_sum.set_column("B:B", 25, format_align_center)
    ws_sum.set_column("C:C", 24, format_align_center)
    ws_sum.set_column("D:D", 20, format_align_center)
    ws_sum.set_column("E:E", 20, format_align_center)
    ws_sum.set_column("F:F", 7, format_align_center)

    # --------------------
    # VULN SUMMARY
    # --------------------
    ws_sum.merge_range("B2:E2", "VULNERABILITY SUMMARY", format_sheet_title_content) # type: ignore
    ws_sum.write("B3", "Threat", format_table_titles)
    ws_sum.write("C3", "Unique Vulns", format_table_titles)
    ws_sum.write("D3", "Hosts affected", format_table_titles)
    ws_sum.write("E3", "Discovered", format_table_titles)

    threat_type_number = len(threat_type_list)
    end_table = 3+threat_type_number
    for i, level in enumerate(threat_type_list, 4):
        ws_sum.write("B{}".format(i), level.capitalize(), format_sheet_title_content)
        ws_sum.write("C{}".format(i), vuln_levels[level], format_align_border)
        ws_sum.write("D{}".format(i), vuln_host_by_level[level], format_align_border)
        ws_sum.write("E{}".format(i), vuln_levels[level] * vuln_host_by_level[level], format_align_border)

    ws_sum.write(f"B{end_table+1}", "Total", format_table_titles)
    ws_sum.write_formula(f"C{end_table+1}", f"=SUM($C$4:$C${end_table})", format_table_titles)
    ws_sum.write_formula(f"D{end_table+1}", f"=SUM($D$4:$D${end_table})", format_table_titles)
    ws_sum.write_formula(f"E{end_table+1}", f"=SUM($E$4:$E${end_table})", format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_summary = workbook.add_chart({'type': 'pie'})
    chart_vulns_summary.add_series({ # type: ignore
        'name': 'vulnerability summary by affected hosts',
        'categories': '={}!B4:B{}'.format(sheet_name, end_table),
        'values': '={}!D4:D{}'.format(sheet_name, end_table),
        'data_labels': {'value': True, 'position': 'outside_end', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
        'points': [
            {'fill': {'color': Config.colors()[each]}} for each in threat_type_list
        ],
    })
    chart_vulns_summary.set_title({'name': 'Vulnerability summary', 'overlay': False, 'name_font': {'name': 'Tahoma'}}) # type: ignore
    chart_vulns_summary.set_size({'width': 500, 'height': 300}) # type: ignore
    chart_vulns_summary.set_legend({'position': 'right', 'font': {'name': 'Tahoma'}}) # type: ignore
    ws_sum.insert_chart("G2", chart_vulns_summary) # type: ignore

    # --------------------
    # VULN BY FAMILY
    # --------------------
    ws_sum.merge_range("B19:C19", "VULNERABILITIES BY FAMILY", format_sheet_title_content) # type: ignore
    ws_sum.write("B20", "Family", format_table_titles)
    ws_sum.write("C20", "Vulnerabilities", format_table_titles)

    last = 21
    for i, (family, number) in enumerate(iter(vuln_by_family.items()), last):
        ws_sum.write("B{}".format(i), family, format_align_border)
        ws_sum.write("C{}".format(i), number, format_align_border)
        last = i

    ws_sum.write("B{}".format(str(last + 1)), "Total", format_table_titles)
    ws_sum.write_formula("C{}".format(str(last + 1)), "=SUM($C$21:$C${})".format(last), format_table_titles)

    # --------------------
    # CHART
    # --------------------
    chart_vulns_by_family = workbook.add_chart({'type': 'pie'})
    chart_vulns_by_family.add_series({ # type: ignore
        'name': 'vulnerability summary by family',
        'categories': '={}!B21:B{}'.format(sheet_name, last),
        'values': '={}!C21:C{}'.format(sheet_name, last),
        'data_labels': {'value': True, 'position': 'best_fit', 'leader_lines': True, 'font': {'name': 'Tahoma'}},
    })
    chart_vulns_by_family.set_title({'name': 'Vulnerabilities by family', 'overlay': False, # type: ignore
                                     'name_font': {'name': 'Tahoma'}})
    chart_vulns_by_family.set_size({'width': 500, 'height': 500}) # type: ignore
    chart_vulns_by_family.set_legend({'position': 'bottom', 'font': {'name': 'Tahoma'}}) # type: ignore
    ws_sum.insert_chart("G19", chart_vulns_by_family) # type: ignore

    # ====================
    # TABLE OF CONTENTS
    # ====================
    sheet_name = "TOC"
    ws_toc = workbook.add_worksheet(sheet_name)
    ws_toc.set_tab_color(Config.colors()['blue'])

    ws_toc.set_column("A:A", 7)
    ws_toc.set_column("B:B", 5)
    ws_toc.set_column("C:C", 70)
    ws_toc.set_column("D:D", 15)
    ws_toc.set_column("E:E", 50)
    ws_toc.set_column("F:F", 7)

    ws_toc.merge_range("B2:E2", "TABLE OF CONTENTS", format_sheet_title_content) # type: ignore
    ws_toc.write("B3", "No.", format_table_titles)
    ws_toc.write("C3", "Vulnerability", format_table_titles)
    ws_toc.write("D3", "CVSS Score", format_table_titles)
    ws_toc.write("E3", "Hosts", format_table_titles)

    # ====================
    # VULN SHEETS
    # ====================
    for i, vuln in enumerate(vuln_info, 1):
        name = sub(r"[\[\]\\\'\"&@#():*?/]", "", vuln.name)
        if len(name) > 27:
            name = "{}..{}".format(name[0:15], name[-10:])
        name = "{:03X}_{}".format(i, name)
        ws_vuln = workbook.add_worksheet(name)
        ws_vuln.set_tab_color(Config.colors()[vuln.level.lower()])

        vuln.version = ""
        if(match := (search(r'Installed version: ((\d|.)+)', vuln.hosts[0][1].result) or search(r'EOL version:( )+((\d|.)+)', vuln.hosts[0][1].result))):
            vuln.version = match.group(1)

        # --------------------
        # TABLE OF CONTENTS
        # --------------------
        ws_toc.write("B{}".format(i + 3), "{:03X}".format(i), format_table_cells)
        ws_toc.write_url("C{}".format(i + 3), "internal:'{}'!A1".format(name), format_table_cells, string=vuln.name)
        ws_toc.write("D{}".format(i + 3), "{:.1f} ({})".format(vuln.cvss, vuln.level.capitalize()),
                     format_toc[vuln.level])
        ws_toc.write("E{}".format(i + 3), "{}".format(', '.join({host.ip for host, _ in vuln.hosts})),
                     format_table_cells)
        ws_vuln.write_url("A1", "internal:'{}'!A{}".format(ws_toc.get_name(), i + 3), format_align_center,
                          string="<< TOC")
        ws_toc.set_row(i + 3, __row_height(name, 150), None)

        # --------------------
        # VULN INFO
        # --------------------
        ws_vuln.set_column("A:A", 7, format_align_center)
        ws_vuln.set_column("B:B", 20, format_align_center)
        ws_vuln.set_column("C:C", 20, format_align_center)
        ws_vuln.set_column("D:D", 50, format_align_center)
        ws_vuln.set_column("E:E", 15, format_align_center)
        ws_vuln.set_column("F:F", 15, format_align_center)
        ws_vuln.set_column("G:G", 20, format_align_center)
        ws_vuln.set_column("H:H", 7, format_align_center)
        content_width = 120

        ws_vuln.write('B2', "Title", format_table_titles)
        ws_vuln.merge_range("C2:G2", vuln.name, format_sheet_title_content) # type: ignore
        ws_vuln.set_row(1, __row_height(vuln.name, content_width), None)

        ws_vuln.write('B3', "Description", format_table_titles)
        ws_vuln.merge_range("C3:G3", vuln.description, format_table_cells) # type: ignore
        ws_vuln.set_row(2, __row_height(vuln.description, content_width), None)

        ws_vuln.write('B4', "Impact", format_table_titles)
        ws_vuln.merge_range("C4:G4", vuln.impact, format_table_cells) # type: ignore
        ws_vuln.set_row(3, __row_height(vuln.impact, content_width), None)

        ws_vuln.write('B5', "Recommendation", format_table_titles)
        ws_vuln.merge_range("C5:G5", vuln.solution, format_table_cells) # type: ignore
        ws_vuln.set_row(4, __row_height(vuln.solution, content_width), None)

        ws_vuln.write('B6', "Version", format_table_titles)
        ws_vuln.merge_range("C6:G6", vuln.version, format_table_cells) # type: ignore
        ws_vuln.set_row(5, __row_height(vuln.version, content_width), None)

        ws_vuln.write('B7', "Details", format_table_titles)
        ws_vuln.merge_range("C7:G7", vuln.insight, format_table_cells) # type: ignore
        ws_vuln.set_row(6, __row_height(vuln.insight, content_width), None)

        ws_vuln.write('B8', "CVEs", format_table_titles)
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"
        ws_vuln.merge_range("C8:G8", cves, format_table_cells) # type: ignore
        ws_vuln.set_row(7, __row_height(cves, content_width), None)

        ws_vuln.write('B9', "CVSS", format_table_titles)
        cvss = float(vuln.cvss)
        if cvss >= 0.0:
            ws_vuln.merge_range("C9:G9", "{:.1f}".format(cvss), format_table_cells) # type: ignore
        else:
            ws_vuln.merge_range("C9:G9", "{}".format("No CVSS"), format_table_cells) # type: ignore

        ws_vuln.write('B10', "Level", format_table_titles)
        ws_vuln.merge_range("C10:G10", vuln.level.capitalize(), format_table_cells) # type: ignore

        ws_vuln.write('B11', "Family", format_table_titles)
        ws_vuln.merge_range("C11:G11", vuln.family, format_table_cells) # type: ignore

        ws_vuln.write('B12', "References", format_table_titles)
        ws_vuln.merge_range("C12:G12", " {}".format(vuln.references), format_table_cells) # type: ignore
        ws_vuln.set_row(10, __row_height(vuln.references, content_width), None)

        ws_vuln.write('C14', "IP", format_table_titles)
        ws_vuln.write('D14', "Host name", format_table_titles)
        ws_vuln.write('E14', "Port number", format_table_titles)
        ws_vuln.write('F14', "Port protocol", format_table_titles)
        ws_vuln.write('G14', "Result", format_table_titles)

        # --------------------
        # AFFECTED HOSTS
        # --------------------
        
        for j, (host, port) in enumerate(vuln.hosts, 15):

            ws_vuln.write("C{}".format(j), host.ip)
            ws_vuln.write("D{}".format(j), host.host_name if host.host_name else "-")

            if port:
                ws_vuln.write("E{}".format(j), "" if port.number == 0 else port.number)
                ws_vuln.write("F{}".format(j), port.protocol)
                ws_vuln.write("G{}".format(j), port.result, format_table_cells)
                ws_vuln.set_row(j, __row_height(port.result, content_width), None)
            else:
                ws_vuln.write("E{}".format(j), "No port info")

    workbook.close()


def export_to_word_by_vuln(vuln_info:list[Vulnerability], threat_type_list:list[str], template=None, output_file:str='openvas_report.docx') -> None:
    """
    Export vulnerabilities info in a Word file.

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param threat_type_list: Type of Threat present
    :type threat_type_list: list(str)

    :param output_file: Filename of the Excel file
    :type output_file: str
    
    :param template: Path to Docx template
    :type template: str

    :raises: TypeError
    """

    import matplotlib.pyplot as plt
    import numpy as np
    import tempfile
    import os

    from docx import Document
    from docx.oxml.shared import qn, OxmlElement
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Cm

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
            
    if not isinstance(threat_type_list, list):
        raise TypeError("Expected list, got '{}' instead".format(type(threat_type_list)))
    else:
        for x in threat_type_list:
            if not isinstance(x, str):
                raise TypeError("Expected str, got '{}' instead".format(type(x)))
            
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        if not isinstance(template, str):
            raise TypeError("Expected str, got '{}' instead".format(type(template)))
    else:
        # == HAMMER PROGRAMMING (beat it into submission) ==
        # I had to use pkg_resources because I couldn't find this template any other way. 
        import pkg_resources
        template = pkg_resources.resource_filename('openvasreporting', 'src/openvas-template.docx')
    
    #mpl_logger = logging.getLogger('matplotlib')
    #mpl_logger.setLevel(logging.NOTSET)
    
    vuln_info, vuln_levels, vuln_host_by_level, vuln_by_family = _get_collections(vuln_info)

    # ====================
    # DOCUMENT PROPERTIES
    # ====================
    document = Document(template)

    doc_prop = document.core_properties
    doc_prop.title = "OpenVAS Report"
    doc_prop.category = "Report"

    document.add_paragraph('OpenVAS Report', style='Title')

    # ====================
    # TABLE OF CONTENTS
    # ====================
    document.add_paragraph('Table of Contents', style='Heading 1')

    par = document.add_paragraph()
    run = par.add_run()
    fld_char = OxmlElement('w:fldChar')  # creates a new element
    fld_char.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instr_text.text = r'TOC \h \z \t "OV-H1toc;1;OV-H2toc;2;OV-H3toc;3;OV-Finding;3"' # type: ignore

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:t')
    fld_char3.text = "# Right-click to update field. #" # type: ignore
    fld_char2.append(fld_char3)

    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)

    document.add_page_break()

    # ====================
    # MANAGEMENT SUMMARY
    # ====================
    document.add_paragraph('Management Summary', style='OV-H1toc')
    document.add_paragraph('< TYPE YOUR MANAGEMENT SUMMARY HERE >')
    document.add_page_break()

    # ====================
    # TECHNICAL FINDINGS
    # ====================
    document.add_paragraph('Technical Findings', style='OV-H1toc')
    document.add_paragraph('The section below discusses the technical findings.')

    # --------------------
    # SUMMARY TABLE
    # --------------------
    document.add_paragraph('Summary', style='OV-H2toc')

    colors_sum = []
    labels_sum = []
    vuln_sum = []
    aff_sum = []

    table_summary = document.add_table(rows=1, cols=3)
    hdr_cells = table_summary.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Risk level').bold = True
    hdr_cells[1].paragraphs[0].add_run('Vulns number').bold = True
    hdr_cells[2].paragraphs[0].add_run('Affected hosts').bold = True

    # Provide data to table and charts
    for level in threat_type_list:
        row_cells = table_summary.add_row().cells
        row_cells[0].text = level.capitalize()
        row_cells[1].text = str(vuln_levels[level])
        row_cells[2].text = str(vuln_host_by_level[level])
        colors_sum.append(Config.colors()[level])
        labels_sum.append(level)
        vuln_sum.append(vuln_levels[level])
        aff_sum.append(vuln_host_by_level[level])

    # --------------------
    # CHART
    # --------------------
    fd, path = tempfile.mkstemp(suffix='.png')

    par_chart = document.add_paragraph()
    run_chart = par_chart.add_run()

    plt.figure()

    pos = np.arange(len(labels_sum))
    width = 0.35

    bars_vuln = plt.bar(pos - width / 2, vuln_sum, width, align='center', label='Vulnerabilities',
                        color=colors_sum, edgecolor='black')
    bars_aff = plt.bar(pos + width / 2, aff_sum, width, align='center', label='Affected hosts',
                       color=colors_sum, edgecolor='black', hatch='//')
    plt.title('Vulnerability summary by risk level')
    plt.subplot().set_xticks(pos)
    plt.subplot().set_xticklabels(labels_sum)
    plt.gca().spines['left'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['bottom'].set_position('zero')
    plt.tick_params(top=False, bottom=True, left=False, right=False,
                    labelleft=False, labelbottom=True)
    plt.subplots_adjust(left=0.0, right=1.0)

    def __label_bars(barcontainer):
        for bar in barcontainer:
            height = bar.get_height()
            plt.gca().text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, str(int(height)),
                           ha='center', color='black', fontsize=9)

    __label_bars(bars_vuln)
    __label_bars(bars_aff)

    plt.legend()

    plt.savefig(path)

    # plt.show()  # DEBUG

    run_chart.add_picture(path, width=Cm(8.0))

    plt.figure()

    values = list(vuln_by_family.values())
    _, _, autotexts = plt.pie(values, labels=vuln_by_family.keys(), autopct='') # type: ignore
    plt.title('Vulnerability by family')
    for i, txt in enumerate(autotexts):
        txt.set_text('{}'.format(values[i]))
    plt.axis('equal')

    plt.savefig(path, bbox_inches='tight')  # bbox_inches fixes labels being cut, however only on save not on show

    # plt.show()  # DEBUG

    run_chart.add_picture(path, width=Cm(8.0))
    os.close(fd)
    os.remove(path)

    # ====================
    # VULN PAGES
    # ====================
    cur_level = ""

    for i, vuln in enumerate(vuln_info, 1):
        # --------------------
        # GENERAL
        # --------------------
        level = vuln.level.lower()

        if level != cur_level:
            document.add_paragraph(
                level.capitalize(), style='OV-H2toc').paragraph_format.page_break_before = True
            cur_level = level
        else:
            document.add_page_break()

        title = "[{}] {}".format(level.upper(), vuln.name)
        document.add_paragraph(title, style='OV-Finding')

        table_vuln = document.add_table(rows=9, cols=3)
        table_vuln.autofit = False

        # COLOR
        # --------------------
        col_cells = table_vuln.columns[0].cells
        col_cells[0].merge(col_cells[7])
        color_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), Config.colors()[vuln.level][1:]))
        col_cells[0]._tc.get_or_add_tcPr().append(color_fill)

        for col_cell in col_cells:
            col_cell.width = Cm(0.42)

        # TABLE HEADERS
        # --------------------
        hdr_cells = table_vuln.columns[1].cells
        hdr_cells[0].paragraphs[0].add_run('Description').bold = True
        hdr_cells[1].paragraphs[0].add_run('Impact').bold = True
        hdr_cells[2].paragraphs[0].add_run('Recommendation').bold = True
        hdr_cells[3].paragraphs[0].add_run('Version').bold = True
        hdr_cells[4].paragraphs[0].add_run('Details').bold = True
        hdr_cells[5].paragraphs[0].add_run('CVSS').bold = True
        hdr_cells[6].paragraphs[0].add_run('CVEs').bold = True
        hdr_cells[7].paragraphs[0].add_run('Family').bold = True
        hdr_cells[8].paragraphs[0].add_run('References').bold = True

        for hdr_cell in hdr_cells:
            hdr_cell.width = Cm(3.58)

        # FIELDS
        # --------------------
        cves = ", ".join(vuln.cves)
        cves = cves.upper() if cves != "" else "No CVE"

        cvss = str(vuln.cvss) if vuln.cvss != -1.0 else "No CVSS"

        vuln.version = ""
        if(match := search(r'Installed version: ((\d|.)+)', vuln.hosts[0][1].result)):
            vuln.version = match.group(1)

        txt_cells = table_vuln.columns[2].cells
        txt_cells[0].text = vuln.description
        txt_cells[1].text = vuln.impact
        txt_cells[2].text = vuln.solution
        txt_cells[3].text = vuln.version
        txt_cells[4].text = vuln.insight
        txt_cells[5].text = cvss
        txt_cells[6].text = cves
        txt_cells[7].text = vuln.family
        txt_cells[8].text = vuln.references

        for txt_cell in txt_cells:
            txt_cell.width = Cm(12.50)

        # VULN HOSTS
        # --------------------
        document.add_paragraph('Vulnerable hosts', style='Heading 4')

        # add coloumn for result per port and resize columns
        table_hosts = document.add_table(cols=5, rows=(len(vuln.hosts) + 1))

        col_cells = table_hosts.columns[1].cells
        for col_cell in col_cells:
            col_cell.width = Cm(3.2)

        col_cells = table_hosts.columns[2].cells
        for col_cell in col_cells:
            col_cell.width = Cm(3.2)

        col_cells = table_hosts.columns[2].cells
        for col_cell in col_cells:
            col_cell.width = Cm(1.6)

        col_cells = table_hosts.columns[3].cells
        for col_cell in col_cells:
            col_cell.width = Cm(1.6)

        col_cells = table_hosts.columns[4].cells
        for col_cell in col_cells:
            col_cell.width = Cm(6.4)

        hdr_cells = table_hosts.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('IP').bold = True
        hdr_cells[1].paragraphs[0].add_run('Host name').bold = True
        hdr_cells[2].paragraphs[0].add_run('Port number').bold = True
        hdr_cells[3].paragraphs[0].add_run('Port protocol').bold = True
        hdr_cells[4].paragraphs[0].add_run('Port result').bold = True

        for j, (host, port) in enumerate(vuln.hosts, 1):
            cells = table_hosts.rows[j].cells
            cells[0].text = host.ip
            cells[1].text = host.host_name if host.host_name else "-"
            if port and port is not None:
                cells[2].text = "-" if port.number == 0 else str(port.number)
                cells[3].text = port.protocol
                cells[4].text = port.result
            else:
                cells[2].text = "No port info"

    document.save(output_file)


def export_to_csv_by_vuln(vuln_info:list[Vulnerability], threat_type_list:list[str], template=None, output_file:str='openvas_report.csv') -> None:
    """
    Export vulnerabilities info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param threat_type_list: Type of Threat present
    :type threat_type_list: list(str)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
            
    if not isinstance(threat_type_list, list):
        raise TypeError("Expected list, got '{}' instead".format(type(threat_type_list)))
    else:
        for x in threat_type_list:
            if not isinstance(x, str):
                raise TypeError("Expected str, got '{}' instead".format(type(x)))
            
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in CSV-output.")

    vuln_info, _, _, _ = _get_collections(vuln_info)

    with open(output_file, 'w') as csvfile:
        fieldnames = ['hostname', 'ip', 'port', 'protocol',
                      'vulnerability', 'cvss', 'threat', 'family',
                      'description', 'detection', 'insight', 'impact', 'affected', 'solution', 'solution_type',
                      'vuln_id', 'cve', 'references']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for vuln in vuln_info:
            for (host, port) in vuln.hosts:
                rowdata = {
                    'hostname': host.host_name,
                    'ip': host.ip,
                    'port': port.number,
                    'protocol': port.protocol,
                    'vulnerability': vuln.name,
                    'cvss': vuln.cvss,
                    'threat': vuln.level,
                    'family': vuln.family,
                    'description': vuln.description,
                    'detection': vuln.detect,
                    'insight': vuln.insight,
                    'impact': vuln.impact,
                    'affected': vuln.affected,
                    'solution': vuln.solution,
                    'solution_type': vuln.solution_type,
                    'vuln_id': vuln.vuln_id,
                    'cve': ' - '.join(vuln.cves),
                    'references': ' - '.join(vuln.references)
                }
                writer.writerow(rowdata)


def export_to_excel_by_host(resulttree: ResultTree, threat_type_list:list[str], template=None, output_file:str='openvas_report.xlsx') -> None:
    """
    Export vulnerabilities info in an Excel file.

    :param resulttree: Vulnerability list info
    :type resulttree: resulttree

    :param threat_type_list: Type of Threat present
    :type threat_type_list: list(str)

    :param template: Not supported in xlsx-output
    :type template: NoneType

    :param output_file: Filename of the Excel file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import xlsxwriter

    if not isinstance(resulttree, ResultTree):
        raise TypeError("Expected ResultTree, got '{}' instead".format(type(resulttree)))
    else:
        for key in resulttree.keys():
            if not isinstance(resulttree[key], Host):
                raise TypeError("Expected Host, got '{}' instead".format(type(resulttree[key])))
            
    if not isinstance(threat_type_list, list):
        raise TypeError("Expected list, got '{}' instead".format(type(threat_type_list)))
    else:
        for x in threat_type_list:
            if not isinstance(x, str):
                raise TypeError("Expected str, got '{}' instead".format(type(x)))
            
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in XSLX-output.")

    # ====================
    # FUNCTIONS
    # ====================
    def __row_height(text, width):
        return (max((len(text) // width), text.count('\n')) + 1) * 15

    workbook = xlsxwriter.Workbook(output_file)

    workbook.set_properties({
        'title': output_file,
        'subject': 'OpenVAS report',
        'author': 'CNPP',
        'category': 'report',
        'keywords': 'OpenVAS, report',
        'comments': 'Forked from : TheGroundZero (https://github.com/TheGroundZero)'})

    # ====================
    # FORMATTING
    # ====================
    workbook.formats[0].set_font_name('Tahoma')

    format_sheet_title_content = workbook.add_format({'font_name': 'Tahoma', 'font_size': 12,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'center', 'valign': 'vcenter', 'border': 1})
    format_table_titles = workbook.add_format({'font_name': 'Tahoma', 'font_size': 11,
                                               'font_color': 'white', 'bold': True,
                                               'align': 'center', 'valign': 'vcenter',
                                               'border': 1,
                                               'bg_color': Config.colors()['blue']})
    format_table_left_item = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                                      'font_color': Config.colors()['blue'], 'bold': True,
                                                      'align': 'left', 'valign': 'vcenter', 'border': 1})
    format_table_cells = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                              'align': 'left', 'valign': 'top',
                                              'border': 1, 'text_wrap': 1})
    format_align_center = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'center', 'valign': 'top'})
    format_align_left = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'left', 'valign': 'top'})
    format_align_right = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'right', 'valign': 'top'})
    format_align_border_left = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'left', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_align_border_right = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'right', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_number_border_right = workbook.add_format({'font_name': 'Tahoma', 'font_size': 10,
                                               'align': 'right', 'valign': 'top',
                                               'border': 1, 'text_wrap': 1})
    format_number_border_right.num_format = '#.00'
    format_toc = {
        'critical': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                         'align': 'center', 'valign': 'top',
                                         'border': 1,
                                         'bg_color': Config.colors()['critical']}),
        'high': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['high']}),
        'medium': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                       'align': 'center', 'valign': 'top',
                                       'border': 1, 'bg_color': Config.colors()['medium']}),
        'low': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                    'align': 'center', 'valign': 'top',
                                    'border': 1, 'bg_color': Config.colors()['low']}),
        'none': workbook.add_format({'font_name': 'Tahoma', 'font_size': 10, 'font_color': 'white',
                                     'align': 'center', 'valign': 'top',
                                     'border': 1, 'bg_color': Config.colors()['none']})
    }

    # ====================
    # SUMMARY SHEET
    # ====================
    sheet_name = "Summary"
    ws_sum = workbook.add_worksheet(sheet_name)
    ws_sum.set_tab_color(Config.colors()['blue'])

    ws_sum.set_column("A:A", 3, format_align_center)
    ws_sum.set_column("B:B", 8, format_align_left)
    ws_sum.set_column("C:C", 30, format_align_left)
    ws_sum.set_column("D:D", 15, format_align_right)
    for pos, _ in enumerate(threat_type_list, 69):
        column = chr(pos)
        ws_sum.set_column("{}:{}".format(column, column), 8, format_align_right)

    threat_len_end = 69 + len(threat_type_list)
    ws_sum.set_column("{}:{}".format(chr(threat_len_end), chr(threat_len_end)), 8, format_align_right) # total
    ws_sum.set_column("{}:{}".format(chr(threat_len_end + 1), chr(threat_len_end + 1)), 7, format_align_center) # severity

    # ---------------------
    # MAX 10 HOSTS 
    # ---------------------
    if len(resulttree) < 10:
        max_hosts = len(resulttree)
    else:
        max_hosts = 10

    # --------------------------
    # HOST SUM SEVERITY SUMMARY
    # --------------------------
    ws_sum.merge_range("B2:{}2".format(chr(threat_len_end + 1)), "Hosts Ranking", format_sheet_title_content) # type: ignore
    ws_sum.write("B3", "#", format_table_titles)
    ws_sum.write("C3", "Hostname", format_table_titles)
    ws_sum.write("D3", "IP", format_table_titles)

    for pos, threat in enumerate(threat_type_list, 69):
        column = chr(pos)
        ws_sum.write("{}3".format(column), threat, format_table_titles)

    ws_sum.write("{}3".format(chr(threat_len_end)), "total", format_table_titles) # total
    ws_sum.write("{}3".format(chr(threat_len_end + 1)), "severity", format_table_titles) # severity
    
    temp_resulttree = resulttree.sorted_keys_by_rank()
    
    for i, key in enumerate(temp_resulttree[:max_hosts], 4):
        ws_sum.write("B{}".format(i), i-3, format_table_left_item)
        ws_sum.write("C{}".format(i), resulttree[key].host_name, format_table_left_item)
        ws_sum.write("D{}".format(i), resulttree[key].ip, format_table_left_item)

        for pos, threat in enumerate(threat_type_list, 69):
            ws_sum.write("{}{}".format(chr(pos), i), resulttree[key].nv[threat], format_align_border_right)

        ws_sum.write("{}{}".format(chr(threat_len_end), i), resulttree[key].nv_total(), format_align_border_right)
        ws_sum.write("{}{}".format(chr(threat_len_end + 1), i), resulttree[key].higher_cvss, 
                     format_toc.get(tmp if (tmp := Config.cvss_level(resulttree[key].higher_cvss)) else ""))

    # --------------------
    # CHART
    # --------------------
    chart_sumcvss_summary = workbook.add_chart({'type': 'column'})

    if not chart_sumcvss_summary:
        raise ValueError(chart_sumcvss_summary)
    
    for pos, threat in enumerate(threat_type_list, 69):
        chart_sumcvss_summary.add_series({
            'name': threat,
            'categories': '={}!D4:D{}'.format(sheet_name, max_hosts + 3),
            'values': '={}!{}4:{}{}'.format(sheet_name, chr(pos), chr(pos), max_hosts + 3),
            'fill': { 'width': 8, 'color': Config.colors()[threat]},
            'border': { 'color': Config.colors()['blue']},
        })


    chart_sumcvss_summary.set_title({'name': 'Hosts by CVSS', 'overlay': False, 'font': {'name': 'Tahoma'}})
    chart_sumcvss_summary.set_size({'width': 750, 'height': 350})
    chart_sumcvss_summary.set_legend({'position': 'left', 'font': {'name': 'Tahoma'}})
    chart_sumcvss_summary.set_x_axis({'label_position': 'bottom'})
    chart_sumcvss_summary.set_x_axis({'num_font': {'name': 'Tahoma', 'size': 8}})
    ws_sum.insert_chart(14, 1, chart_sumcvss_summary)

    # ====================
    # TABLE OF CONTENTS
    # ====================
    sheet_name = "TOC"
    ws_toc = workbook.add_worksheet(sheet_name)
    ws_toc.set_tab_color(Config.colors()['blue'])

    ws_toc.set_column("A:A", 3, format_align_center)
    ws_toc.set_column("B:B", 8, format_align_left)
    ws_toc.set_column("C:C", 30, format_align_left)
    ws_toc.set_column("D:D", 15, format_align_right)

    for pos, _ in enumerate(threat_type_list, 69):
        column = chr(pos)
        ws_toc.set_column("{}:{}".format(column, column), 8, format_align_right)

    threat_len_end = 69 + len(threat_type_list)
    ws_toc.set_column("{}:{}".format(chr(threat_len_end), chr(threat_len_end)), 8, format_align_right) # total
    ws_toc.set_column("{}:{}".format(chr(threat_len_end + 1), chr(threat_len_end + 1)), 7, format_align_center) # severity

    # --------------------------
    # HOST SUM SEVERITY SUMMARY
    # --------------------------
    ws_toc.merge_range("B2:{}2".format(chr(threat_len_end + 1)), "Hosts Ranking", format_sheet_title_content) # type: ignore
    ws_toc.write("B3", "#", format_table_titles)
    ws_toc.write("C3", "Hostname", format_table_titles)
    ws_toc.write("D3", "IP", format_table_titles)

    for pos, threat in enumerate(threat_type_list, 69):
        column = chr(pos)
        ws_toc.write("{}3".format(column), threat, format_table_titles)

    ws_toc.write("{}3".format(chr(threat_len_end)), "total", format_table_titles) # total
    ws_toc.write("{}3".format(chr(threat_len_end + 1)), "severity", format_table_titles) # severity
    
    # ====================
    # HOST SHEETS
    # ====================
    for i, key in enumerate(temp_resulttree, 1):
 
        # this host has any vulnerability whose cvss severity >= min_level?
        if len(resulttree[key].vuln_list) == 0:
            continue

        name = "{:03X} - {}".format(i, resulttree[key].ip)
        ws_host = workbook.add_worksheet(name)
        ws_host.set_tab_color(Config.cvss_color(resulttree[key].higher_cvss))
        ws_host.write_url("A1", "internal:'{}'!A{}".format(ws_toc.get_name(), i + 3), format_align_center,
                          string="<< TOC")

        # --------------------
        # TABLE OF CONTENTS
        # --------------------
        ws_toc.write("B{}".format(i + 3), "{:03X}".format(i), format_table_cells)
        ws_toc.write_url("C{}".format(i + 3), "internal:'{}'!A1".format(name), format_table_cells,
                         string=resulttree[key].host_name)
        ws_toc.write("D{}".format(i+3), resulttree[key].ip, format_align_border_left)

        for pos, threat in enumerate(threat_type_list, 69):
            column = chr(pos)
            ws_toc.write("{}{}".format(column, i+3), resulttree[key].nv[threat], format_align_border_right)

        ws_toc.write("{}{}".format(chr(threat_len_end), i+3), resulttree[key].nv_total(), format_align_border_right)
        ws_toc.write("{}{}".format(chr(threat_len_end + 1), i+3), resulttree[key].higher_cvss, 
                     format_toc.get(tmp if (tmp := Config.cvss_level(resulttree[key].higher_cvss)) else ""))
        ws_toc.set_row(i + 3, __row_height(name, 150), None)

        # --------------------
        # HOST VULN LIST
        # --------------------
        ws_host.set_column("A:A", 7, format_align_center)
        ws_host.set_column("B:B", 12, format_align_center) # cvss - (level)
        ws_host.set_column("C:C", 22, format_align_center) # name
        ws_host.set_column("D:D", 22, format_align_center) # version
        ws_host.set_column("E:E", 22, format_align_center) # oid
        ws_host.set_column("F:F", 10, format_align_center) # port.port/port.num
        ws_host.set_column("G:G", 22, format_align_center) # family
        ws_host.set_column("H:H", 22, format_align_center) # description
        ws_host.set_column("I:I", 12, format_align_center) # recomendation (solution)
        ws_host.set_column("J:J", 12, format_align_center) # recomendation type (solution_type)
        ws_host.set_column("K:K", 7, format_align_center)
        
        ws_host.merge_range("B2:K2", resulttree[key].ip + ' - ' + resulttree[key].host_name, format_sheet_title_content) # type: ignore
        ws_host.write('B3', "CVSS", format_table_titles)
        ws_host.write('C3', "Name", format_table_titles)
        ws_host.write('D3', "Version", format_table_titles)
        ws_host.write('E3', "oid", format_table_titles)
        ws_host.write('F3', "Port", format_table_titles)
        ws_host.write('G3', "Family", format_table_titles)
        ws_host.write('H3', "Description", format_table_titles)
        ws_host.write('I3', "Recomendation", format_table_titles)
        ws_host.write('J3', "Type of fix", format_table_titles)


        for j, vuln in enumerate(resulttree[key].vuln_list, 4):
            ws_host.write('B{}'.format(j), "{:.2f} ({})".format(vuln.cvss, vuln.level),
                          format_toc[vuln.level])
            ws_host.write('C{}'.format(j), vuln.name, format_align_border_left)
            ws_host.write('D{}'.format(j), vuln.version, format_align_border_left)
            ws_host.write('E{}'.format(j), vuln.vuln_id, format_align_border_left)
            port = vuln.hosts[0][1]
            if port is None or port.number == 0:
                portnum = 'general' 
            else: 
                portnum = str(port.number)
            ws_host.write('F{}'.format(j), portnum + '/' + port.protocol, format_align_border_left)
            ws_host.write('G{}'.format(j), vuln.family, format_align_border_left)
            ws_host.write('H{}'.format(j), vuln.description.replace('\n', ' '), format_align_border_left)
            ws_host.write('I{}'.format(j), vuln.solution.replace('\n', ' '), format_align_border_left)
            ws_host.write('J{}'.format(j), vuln.solution_type, format_align_border_left)
            max_len = max(len(vuln.name), len(vuln.description), len(vuln.solution))
            ws_host.set_row(j-1, (int(max_len/30)+1)*15)
        
    workbook.close()


def export_to_csv_by_host(resulttree: ResultTree, threat_type_list:list[str], template=None, output_file:str='openvas_report.csv') -> None:
    """
    Export vulnerabilities info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param threat_type_list: Type of Threat present
    :type threat_type_list: list(str)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv
    
    if not isinstance(resulttree, ResultTree):
        raise TypeError("Expected ResultTree, got '{}' instead".format(type(resulttree)))
    else:
        for x in resulttree.values():
            if not isinstance(x, Host):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    
    if not isinstance(threat_type_list, list):
        raise TypeError("Expected list, got '{}' instead".format(type(threat_type_list)))
    else:
        for x in threat_type_list:
            if not isinstance(x, str):
                raise TypeError("Expected str, got '{}' instead".format(type(x)))
            
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
    if template is not None:
        raise NotImplementedError("Use of template is not supported in CSV-output.")
  
    sortedresults = resulttree.sortedbysumcvss()

    with open(output_file, 'w') as csvfile:
        fieldnames = ['hostname', 'ip', 'port', 'protocol',
                      'vulnerability', 'cvss', 'threat', 'family',
                      'description', 'detection', 'version', 'insight', 'impact', 'affected', 'solution', 'solution_type',
                      'vuln_id', 'cve', 'references']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for key in sortedresults:
            for vuln in resulttree[key].vuln_list:
                if Config.cvss_level(vuln.cvss) in threat_type_list:
                    rowdata = {
                        'hostname': resulttree[key].host_name,
                        'ip': resulttree[key].ip,
                        'port': vuln.hosts[0][1].number,
                        'protocol': vuln.hosts[0][1].protocol,
                        'vulnerability': vuln.name,
                        'cvss': vuln.cvss,
                        'threat': vuln.level,
                        'family': vuln.family,
                        'description': vuln.description,
                        'detection': vuln.detect,
                        'version': vuln.version,
                        'insight': vuln.insight,
                        'impact': vuln.impact,
                        'affected': vuln.affected,
                        'solution': vuln.solution,
                        'solution_type': vuln.solution_type,
                        'vuln_id': vuln.vuln_id,
                        'cve': ' - '.join(vuln.cves),
                        'references': ' - '.join(vuln.references) if isinstance(vuln.references, list) else vuln.references
                    }
                    writer.writerow(rowdata)


def export_summary_to_csv(vuln_info:list[Vulnerability], threat_type_list:list[str], template=None, output_file:str='openvas_summary_report.csv') -> None:
    """
    Export summary info in a Comma Separated Values (csv) file

    :param vuln_info: Vulnerability list info
    :type vuln_info: list(Vulnerability)

    :param threat_type_list: Type of Threat present
    :type threat_type_list: list(str)

    :param template: Not supported in csv-output
    :type template: NoneType

    :param output_file: Filename of the csv file
    :type output_file: str

    :raises: TypeError, NotImplementedError
    """

    import csv

    if not isinstance(vuln_info, list):
        raise TypeError("Expected list, got '{}' instead".format(type(vuln_info)))
    else:
        for x in vuln_info:
            if not isinstance(x, Vulnerability):
                raise TypeError("Expected Vulnerability, got '{}' instead".format(type(x)))
    
    if not isinstance(threat_type_list, list):
        raise TypeError("Expected list, got '{}' instead".format(type(threat_type_list)))
    else:
        for x in threat_type_list:
            if not isinstance(x, str):
                raise TypeError("Expected str, got '{}' instead".format(type(x)))
            
    if not isinstance(output_file, str):
        raise TypeError("Expected str, got '{}' instead".format(type(output_file)))
    else:
        if not output_file:
            raise ValueError("output_file must have a valid name.")
        
    if template is not None:
        raise NotImplementedError("Use of template is not supported in CSV-output.")

    vuln_info, vuln_levels, vuln_host_by_level, _ = _get_collections(vuln_info)

    with open(output_file, 'w') as csvfile:
        fieldnames = ['level', 'count', 'host_count']
        writer = csv.DictWriter(csvfile, dialect='excel', fieldnames=fieldnames)
        writer.writeheader()

        for _, level in enumerate(Config.levels().values(), 4):
            rowdata = {
                'level': level,
                'count': vuln_levels[level],
                'host_count': vuln_host_by_level[level]
            }
            writer.writerow(rowdata)
